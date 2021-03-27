from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn,nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
import docx
import datetime
import re
from num2words import num2words
import os
import time
import copy
from docx.enum.text import WD_COLOR_INDEX



def suffix(d):
            return 'th' if 11<=d<=13 else {1:'st',2:'nd',3:'rd'}.get(d%10, 'th')

def custom_strftime(format, t):
    return t.strftime(format).replace('{S}', str(t.day) + suffix(t.day))


class createDoc():
    def __init__(self):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        style = self.document.styles['Normal']
        style.paragraph_format.space_after = 0
        style.paragraph_format.space_before = 0
        style.paragraph_format.line_spacing = 1.115
        font = style.font
        font.name = 'Book Antiqua'
        font.size = docx.shared.Pt(11)

    def formatMinutes(self,MeetingNo,fy):
        section = self.document.sections[0]
        header = section.header
        border = header.add_table(1,1,Cm(5.11))
        border.style='Table Grid'
        border.alignment = WD_TABLE_ALIGNMENT.CENTER
        border.columns[0].cells[0].height = Cm(1)
        self.headline = border.rows[0].cells[0].paragraphs[0]
        self.headline.paragraph_format.space_before=0
        self.headline.alignment = 1
        self.headline.paragraph_format.space_after = 0
        self.headline.paragraph_format.line_spacing = 1.125
        border.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        headText = self.headline.add_run('MINUTES BOOK')
        headText.font.size = docx.shared.Pt(16)
        serialNo = header.add_paragraph(f'Serial No.:{MeetingNo}/{fy}')
        serialNo.alignment = 2
        sec_pr = self.document.sections[0]._sectPr
        ignoreheader = OxmlElement('w:bordersDoNotSurroundHeader')
        ignorefooter = OxmlElement('w:bordersDoNotSurroundFooter')
        pg_borders = OxmlElement('w:pgBorders')
        pg_borders.set(qn('w:offsetFrom'), 'text')
        for border_name in ('top', 'left', 'bottom', 'right',): 
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'single') 
            border_el.set(qn('w:sz'), '4') 
            if border_name in ('top','bottom'):
                border_el.set(qn('w:space'), '1')
            else:
                border_el.set(qn('w:space'), '4')
            border_el.set(qn('w:color'), 'auto')
            pg_borders.append(border_el) 
        sec_pr.append(pg_borders)
        sec_pr.append(ignoreheader)
        sec_pr.append(ignorefooter)
        footer = section.footer
        pageNo = footer.add_paragraph('Page No.:')
        pageNo.alignment = 1
        pageNo.paragraph_format.space_after = 0
        chairIni = footer.add_paragraph("Chairman's Initial")
        chairIni.alignment = 2
        chairIni.paragraph_format.space_after = 0

    def Boardmeetingdetails(self,companyName,MeentingNo,meetingtype,MeetingName,financialYear,TMZ,DayDate,Time, Venue,isTable=False):
        self.document.add_paragraph()
        CoName = self.document.add_paragraph()
        CoName.alignment = 1
        coNameRun = CoName.add_run(companyName.title())
        coNameRun.font.bold = True
        coNameRun.font.size = docx.shared.Pt(12)
        self.document.add_paragraph()
        MeetNo = self.document.add_paragraph()
        MeetNo.alignment = 1
        if not MeentingNo=='':
            if MeetingName=='Board Meeting' or meetingtype=='Committee Meeting':
                MeentingNo = num2words(int(MeentingNo),True)
            else:
                MeentingNo = ' THE PROCEEDINGS OF'
        MeetNoRun = MeetNo.add_run(f"MINUTES OF {MeentingNo.upper()} {meetingtype.upper()}")
        MeetNoRun.font.bold = True
        self.document.add_paragraph()
        DayDate = custom_strftime('%A, {S} %B, %Y', DayDate)
        Time = Time.strftime(f'%I:%M %p ({TMZ}) ')
        textDict = {'Financial Year':financialYear,'Day & Date':DayDate,'Time':Time,'Venue':Venue}
        if isTable:
            infoTable = self.document.add_table(4,3,style='Table Grid')
            infoTable.alignment = WD_TABLE_ALIGNMENT.CENTER            
            for row,key in enumerate(textDict.keys()):
                Text1 = infoTable.rows[row].cells[0].paragraphs[0]
                infoTable.rows[row].cells[0].width = Cm(3.56)
                Text1.add_run(key).bold=True
                Text2 = infoTable.rows[row].cells[1].paragraphs[0]
                Text2.add_run(':').bold=True
                infoTable.rows[row].cells[1].width = Cm(0.73)
                Text3 = infoTable.rows[row].cells[2].paragraphs[0]
                self.splitbysup(Text3,textDict[key],True)
                infoTable.rows[row].cells[2].width = Cm(7.14)
        else:
            for key in textDict.keys():
                Text = key+': '+ textDict[key]
                paragraph = self.document.add_paragraph()
                self.splitbysup(paragraph,Text,True)
            

    def BMattendance(self,attendeesList, MeetingType):
        self.document.add_paragraph()
        CoName = self.document.add_paragraph()
        if 'general' in MeetingType.lower():
            CoName.alignment = 1
            coNameRun = CoName.add_run('Present')
            coNameRun.font.bold = True

        else:
            CoName.alignment = 0
            coNameRun = CoName.add_run('Members Present')
            coNameRun.font.bold = True
            MeetNo = self.document.add_paragraph()
            MeetNo.alignment = 0
            MeetNoRun = MeetNo.add_run('In Person:')
            MeetNoRun.font.bold = True
        self.document.add_paragraph()
        maxstrLen = 0
        for attendee in attendeesList:
##            maxstrLen = maxstrLen if maxstrLen>len(attendee[0]) else len(attendee)
##            ntab = int((maxstrLen/6)+4)
            dirPara = self.document.add_paragraph()
            if 'general' in MeetingType.lower():
                dirPara.add_run(attendee[0]+', '+attendee[1])
            else:  
                tab_stops = dirPara.paragraph_format.tab_stops
                tab_stop = tab_stops.add_tab_stop(Cm(10))
                dirPara.add_run(attendee[0]+'\t\t: '+attendee[1])
            
    def Agendas(self,commencTime,Agendas,signplace,signdate):
        self.document.add_paragraph()
        self.document.add_paragraph(f"The meeting commenced at {commencTime}.")
        self.document.add_paragraph()
        AgendaCounter = 1
        for agenda in Agendas:
            Title = agenda['TITLE'].upper()
            titlePara = self.document.add_paragraph()
            self.splitbysup(titlePara,Title,True)
            titlePara.alignment = 3
            self.document.add_paragraph()
            if agenda['NARRATION']!='':
                for narration in agenda['NARRATION']:
                    narrationPara = self.document.add_paragraph()
                    narrationPara.alignment = 3
                    self.highlightBlanks(narrationPara,narration,isbold = False)
                    self.document.add_paragraph()
            if agenda['RESOLUTION']!='':
                for resolution in agenda['RESOLUTION']:
                    resolutionPara = self.document.add_paragraph()
                    resolutionPara.alignment = 3
                    self.formatResolution(resolutionPara,resolution)
                    self.document.add_paragraph()
        self.document.add_paragraph('')
        PlacePara = self.document.add_paragraph()
        PlacePara.add_run('Place\t: ').bold = True
        PlacePara.add_run(signplace)
        datePara = self.document.add_paragraph()
        datePara.add_run('Date\t: ').bold = True
        datePara.add_run(signdate)
        datePara.add_run('\t'*8+'CHAIRMAN').bold = True
        
                    

                
        
    def formatResolution(self,paragraph,text):
        splitTxt =  re.split(r'(RESOLVED THAT|RESOLVED FURTHER THAT|resolved that|resolved further that)',text)
        for x in splitTxt:
            if x=='':
                pass
            elif x.lower() in ('resolved that','resolved further that'):
                text = paragraph.add_run(x)
                text.font.bold = True
            else:
                text = self.highlightBlanks(paragraph,x)
                
        
        
        
    def splitbysup(self,paragraph,text, isbold = False):
       splitTxt =  re.split(r'(\d(Th|TH|ST|St|th|st|nd|ND|Nd|rd|Rd|RD))',text)
       for x in splitTxt:
                if not(re.match('\d(Th|TH|th|st|ST|St|nd|ND|Nd|rd|Rd|RD)', x))==None:
                    text = paragraph.add_run(re.match('\d',x).group())
                elif x=='':
                    pass
                elif x.lower() in ('th','st','nd','rd'):
                    text = paragraph.add_run(x)
                    text.font.superscript = True
                else:
                    text = paragraph.add_run(x)
                if isbold:
                    text.font.bold = True

    def highlightBlanks(self,paragraph,Text,isbold = False):
        splitTxt =  re.split(r'(_){3,7}',Text)
        for x in splitTxt:
            print(x)
            if x=='':
                pass
            elif '_' in x:
                text = paragraph.add_run(x*5)
                text.font.highlight_color = WD_COLOR_INDEX.YELLOW
                if isbold:
                    text.font.bold = True
            else:
                text = self.splitbysup(paragraph,x,isbold = isbold)

    def saveDoc(self,filename):
        xfile=os.path.splitext(filename)
        if xfile[1]!='docx':
            filename = xfile[0]+'.docx'
        self.document.save(filename)
