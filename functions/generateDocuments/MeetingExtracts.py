from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn,nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
import docx
import datetime
import re
from num2words import num2words
import os
import time
from docx.enum.text import WD_COLOR_INDEX

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def suffix(d):
            return 'th' if 11<=d<=13 else {1:'st',2:'nd',3:'rd'}.get(d%10, 'th')

def custom_strftime(format, t):
    return t.strftime(format).replace('{S}', str(t.day) + suffix(t.day))

class createDoc():
    def __init__(self):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        style = self.document.styles['Normal']
        style.paragraph_format.space_after = 0
        style.paragraph_format.space_before = 0
        style.paragraph_format.line_spacing = 1.115
        font = style.font
        font.name = 'Book Antiqua'
        font.size = docx.shared.Pt(11)
    
    def letterhead(self,companyName, reg_off,cin, email,phone, NoLetterHead = False):
        if not NoLetterHead:
            self.title = self.document.add_paragraph()
            self.title.alignment = 1
            self.title.paragraph_format.space_after = docx.shared.Pt(0)
            self.addressblock = self.document.add_paragraph()
            self.addressblock.paragraph_format.line_spacing = 1
            self.addressblock.alignment = 1
            self.addressblock.paragraph_format.space_after = 0
            self.addressblock.paragraph_format.space_before = 0
            COMP_NAME = self.title.add_run(companyName)
            COMP_NAME.bold = True
            COMP_NAME.font.size = docx.shared.Pt(14)
            REG_ADD=self.addressblock.add_run(f'Regd. Office: ')
            self.splitbysup(self.addressblock,reg_off,fontSize=10)
            REG_ADD.font.size = docx.shared.Pt(10)
            self.addressblock = self.document.add_paragraph()
            self.addressblock.paragraph_format.line_spacing = 1
            self.addressblock.alignment = 1
            self.addressblock.paragraph_format.space_after = 0
            self.addressblock.paragraph_format.space_before = 0
            CIN=self.addressblock.add_run(f'CIN: {cin}, ')
            CIN.font.size = docx.shared.Pt(10)
            emailid = self.addressblock.add_run(f'Email Id.: {email}, ')
            emailid.underline = True
            emailid.font.color.rgb = RGBColor(0, 0, 255)
            emailid.font.size = docx.shared.Pt(10)
            telno = self.addressblock.add_run(f'Tel. No.: {phone}')
            telno.font.size = docx.shared.Pt(10)
            insertHR(self.addressblock)
        else:
            NoHead = self.document.add_paragraph("[ON THE LETTERHEAD OF THE COMPANY]")
            NoHead.alignment=1
            self.document.add_paragraph()
            self.document.add_paragraph()

    def ExtractHeader(self,MeetingType,companyName, address, date, time,finanYear = None, BMNo = None,):
        BMNum = ''
        ForFY = ''
        if BMNo!=None and finanYear!=None:
            self.MeetNo = self.document.add_paragraph()
            self.highlightBlanks(self.MeetNo,f"Serial No.:{BMNo}/{finanYear}",True)
            if str(BMNo).isnumeric():
                BMNum = num2words(int(BMNo),True)+" "
            else:
                BMNum = '__________'
            ForFY = f"FOR THE FINANCIAL YEAR {finanYear} "
        self.document.add_paragraph()
        self.extractHeaderblock = self.document.add_paragraph()
        self.extractHeaderblock.alignment=3
        #mdate =  datetime.datetime.strptime(date,'%d/%m/%Y')
        time = time.strftime('%I:%M %p (IST)')
        ExtractHeader = f"EXTRACTS OF MINUTES OF THE {BMNum}{MeetingType.upper()} {ForFY}OF {companyName.upper()} HELD ON {custom_strftime('%A, {S} %B, %Y', date).upper()} AT {time.upper()} AT {address.upper()}."
        self.highlightBlanks(self.extractHeaderblock,ExtractHeader.upper(),isbold = True)

    def ExtractText (self,title='', Narration='', Resolution=''):
        self.document.add_paragraph()
        self.titleblock = self.document.add_paragraph()
        self.titleblock.alignment=3
        TITLE = self.titleblock.add_run(title.upper())
        TITLE.bold = True
        TITLE.underline = True
        TITLE.alignment = 3
        self.document.add_paragraph()
        if Narration:
            for narration in Narration:
                if narration.strip()!='':
                    narrationPara = self.document.add_paragraph()
                    narrationPara.alignment = 3
                    self.highlightBlanks(narrationPara,narration,isbold = False)
                    self.document.add_paragraph()
        if Resolution:
            for resolution in Resolution:
                if resolution.strip()!='':
                    resolutionPara = self.document.add_paragraph()
                    resolutionPara.alignment = 3
                    self.formatResolution(resolutionPara,resolution)
                    self.document.add_paragraph()

    def Signatures(self, CompanyName, Name, Designation, Din, Address):
        forBlock = self.document.add_paragraph('\n\n')
        forBlock.alignment = 1
        forBlock.add_run('//CERTIFIED TRUE COPY//\n').bold=True
        forBlock.add_run('for ')
        forBlock.add_run(CompanyName).bold=True
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)
        NameBlock = self.document.add_paragraph(Name.title())
        NameBlock.paragraph_format.line_spacing = 1
        NameBlock.paragraph_format.space_after = Pt(0)
        NameBlock.alignment =1
        DesignationBlock=self.document.add_paragraph("("+Designation+", "+"DIN:"+Din+")")
        DesignationBlock.paragraph_format.line_spacing = 1
        DesignationBlock.alignment =1
        DesignationBlock.paragraph_format.space_after = Pt(0)
        AddressList = Address.split(',')
        AddressString = ''
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<40 and len(AddressString+', '+AddressList[x].strip())<40:
                AddressString = AddressString+', '+AddressList[x].strip()
            else:
                AddressM = self.document.add_paragraph()
                self.splitbysup(AddressM,AddressString)
                AddressM.paragraph_format.line_spacing = 1
                AddressM.alignment = 1
                AddressM.paragraph_format.space_after = Pt(0)
                AddressString = ''
                AddressString = AddressString+AddressList[x].strip()
        if AddressString!='':
            AddressM = self.document.add_paragraph()
            self.splitbysup(AddressM,AddressString)
            AddressM.paragraph_format.line_spacing = 1
            AddressM.alignment = 1
            AddressM.paragraph_format.space_after = Pt(0)
            

    def NoticeSignatures(self, CompanyName, Name, Designation, Din, Address):
        self.document.add_paragraph()
        self.document.add_paragraph()
        forBlock = self.document.add_paragraph()
        forBlock.alignment = 1
        forBlock.add_run('for ')
        forBlock.add_run(CompanyName).bold=True
        SignatureBlock = self.document.add_paragraph()
        self.document.add_paragraph()
        NameBlock = self.document.add_paragraph(Name)
        NameBlock.paragraph_format.line_spacing = 1
        NameBlock.paragraph_format.space_after = Pt(0)
        NameBlock.alignment =1
        DesignationBlock=self.document.add_paragraph("("+Designation+", "+"DIN:"+Din+")")
        DesignationBlock.paragraph_format.line_spacing = 1
        DesignationBlock.alignment =1
        DesignationBlock.paragraph_format.space_after = Pt(0)
        AddressList = Address.split(',')
        AddressString = ''
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<32 and len(AddressString+', '+AddressList[x].strip())<32:
                AddressString = AddressString+', '+AddressList[x].strip()
            else:
                AddressM = self.document.add_paragraph(AddressString)
                AddressM.paragraph_format.line_spacing = 1
                AddressM.alignment = 1
                AddressM.paragraph_format.space_after = Pt(0)
                AddressString = ''
                AddressString = AddressString+AddressList[x].strip()
        if AddressString!='':
            AddressM = self.document.add_paragraph(AddressString)
            AddressM.paragraph_format.line_spacing = 1
            AddressM.alignment = 1
            AddressM.paragraph_format.space_after = Pt(0)


    def formatResolution(self,paragraph,text):
        splitTxt =  re.split(r'(RESOLVED THAT|RESOLVED FURTHER THAT|resolved that|resolved further that)',text)
        for x in splitTxt:
            if x=='':
                pass
            elif x.lower() in ('resolved that','resolved further that'):
                text = paragraph.add_run(x)
                text.font.bold = True
            else:
                text = self.highlightBlanks(paragraph,x)
                
        
        
        
    def splitbysup(self,paragraph,text,fontSize = 11,isbold = False):
       splitTxt =  re.split(r'(\d(TH|Th|ST|St|th|st|nd|Nd|ND|rd|Rd|RD))',text)
       for x in splitTxt:
                if not(re.match('\d(TH|Th|ST|St|th|st|nd|Nd|ND|rd|Rd|RD)', x))==None:
                    text = paragraph.add_run(re.match('\d',x).group())
                    text.font.size = docx.shared.Pt(fontSize)
                elif x=='':
                    pass
                elif x.lower() in ('th','st','nd','rd'):
                    text = paragraph.add_run(x)
                    text.font.superscript = True
                    text.font.size = docx.shared.Pt(fontSize)
                else:
                    text = paragraph.add_run(x)
                    text.font.size = docx.shared.Pt(fontSize)
                if isbold:
                    text.font.bold = True
                    
    def highlightBlanks(self,paragraph,Text,isbold = False):
        splitTxt =  re.split(r'(_){3,}',Text)
        for x in splitTxt:
            print(x)
            if x=='':
                pass
            elif '_' in x:
                text = paragraph.add_run(x*5)
                text.font.highlight_color = WD_COLOR_INDEX.YELLOW
                if isbold:
                    text.font.bold = True
            else:
                text = self.splitbysup(paragraph,x,isbold = isbold)

            
        
    def saveDoc(self,filename):
        xfile=os.path.splitext(filename)
        if xfile[1]!='docx':
            filename = xfile[0]+'.docx'
        self.document.save(filename)
