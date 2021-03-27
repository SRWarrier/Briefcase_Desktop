import docx
from docx.shared import Cm
from functions import currency 
import math
import os
from docx.enum.table import WD_TABLE_ALIGNMENT
from babel.numbers import format_decimal

class createDoc():
    def __init__(self):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(0.25)
            section.left_margin = Cm(2.7)
            section.right_margin = Cm(2.7)
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Book Antiqua'
        font.size = docx.shared.Pt(11)
        header = self.document.sections[0].header
        headPara = header.add_paragraph()
        headPara.paragraph_format.left_indent = -Cm(0.25)
        headPara.add_run().add_picture(r'functions/generateDocuments/header.png',width=Cm(16.03), height=Cm(1.53))
        footer = self.document.sections[0].footer
        footPara = footer.add_paragraph()
        footPara.paragraph_format.left_indent = -Cm(3.0)
        footPara.add_run().add_picture(r'functions/generateDocuments/footer.png',width=Cm(21.03), height=Cm(2.04))
        self.document.sections[0].footer_distance = Cm(0)
        self.bill = self.document.add_table(5,4,"Table Grid")
        for col in range(4):
            if col==0 or col==1:
                cellwid = 5.81
            elif col==2:
                cellwid = 2.5
            else:
                cellwid = 3.25
            for cell in self.bill.columns[col].cells:
                cell.width = Cm(cellwid)
        a = self.bill.cell(0, 0)
        b = self.bill.cell(0, 1)
        self.AddressBlock = a.merge(b)
        a = self.bill.cell(0, 2)
        b = self.bill.cell(1, 2)
        self.BillBlock = a.merge(b)
        a = self.bill.cell(0, 3)
        b = self.bill.cell(1, 3)
        self.BillValueBlock = a.merge(b)
        a = self.bill.cell(2, 0)
        b = self.bill.cell(2, 1)
        self.ParticularsTitleBlock = a.merge(b)
        a = self.bill.cell(3, 0)
        b = self.bill.cell(3, 1)
        self.ParticularsBlock = a.merge(b)
        a = self.bill.cell(4, 0)
        b = self.bill.cell(4, 1)
        self.TotalBlock = a.merge(b)


    def BillDetails(self, ClientName, ClientAddress,billNo, BillDate, PONo, GSTIN, StateCode):
        self.bill.rows[0].cells[0].paragraphs[0].add_run("To,")
        self.bill.rows[0].cells[0].add_paragraph().add_run(ClientName).bold = True
        self.bill.rows[0].cells[0].add_paragraph().add_run(ClientAddress)
        self.bill.rows[1].cells[0].paragraphs[0].add_run(f"GSTN:{GSTIN}")
        self.bill.rows[1].cells[1].paragraphs[0].add_run(f"State Code: :{StateCode}")
        Para1 = self.bill.rows[0].cells[2].paragraphs[0]
        Para1.alignment = 2
        Para1.add_run("Bill No.:")
        Para2 = self.bill.rows[0].cells[2].add_paragraph()
        Para2.alignment = 2
        Para2.add_run("Date:")
        Para3 = self.bill.rows[0].cells[2].add_paragraph()
        Para3.alignment = 2
        Para3.add_run("PO No:")
        self.bill.rows[0].cells[3].paragraphs[0].add_run(billNo)
        self.bill.rows[0].cells[3].add_paragraph().add_run(BillDate)
        self.bill.rows[0].cells[3].add_paragraph().add_run(PONo)
        Part4 = self.bill.rows[2].cells[0].paragraphs[0]
        Part4.alignment = 1
        Part4.add_run("Particulars").bold = True
        Part5 = self.bill.rows[2].cells[2].paragraphs[0]
        Part5.alignment = 1
        Part5.add_run("SAC/HSN").bold = True
        Part6 = self.bill.rows[2].cells[2].add_paragraph()
        Part6.alignment = 1
        Part6.add_run("Code:").bold = True
        Part7 = self.bill.rows[2].cells[3].paragraphs[0]
        Part7.alignment = 1
        Part7.add_run("Rs.").bold = True

    def BillParticulars(self,particularsList, GstRates, TotalAmount,  isRegular = True):
        if isRegular:
            Part4 = self.bill.rows[3].cells[0].paragraphs[0]
            Part4.alignment = 3
            Part4.add_run("Please arrange to release payment for the following Professional Services rendered to your organization.")
            Part5 = self.bill.rows[3].cells[0].add_paragraph()
            self.bill.rows[3].cells[0].paragraphs[0]
            self.bill.rows[3].cells[3].paragraphs[0]
            self.bill.rows[3].cells[3].add_paragraph()
            self.bill.rows[3].cells[3].add_paragraph()
            self.bill.rows[3].cells[2].paragraphs[0]
            self.bill.rows[3].cells[2].add_paragraph()
            self.bill.rows[3].cells[2].add_paragraph()
        for item in  particularsList:
            Part6 = self.bill.rows[3].cells[0].add_paragraph(item[0], style='List Bullet')
            Part6.paragraph_format.left_indent = Cm(1.5)
            Part6.alignment = 3
            if len(item[0])/45>0:
                for x in range(int(len(item[0])/45)):
                    self.bill.rows[3].cells[3].add_paragraph()
                    if x!=0:
                        self.bill.rows[3].cells[2].add_paragraph()
            else:
                self.bill.rows[3].cells[3].add_paragraph()
                self.bill.rows[3].cells[3].add_paragraph()
                self.bill.rows[3].cells[3].add_paragraph()
                self.bill.rows[3].cells[2].add_paragraph()
                self.bill.rows[3].cells[2].add_paragraph()
            Amount = format_decimal(float(item[2]), locale='en_IN')
            deciNum =  Amount if Amount.find('.')!=-1 else Amount+'.00'
            Part6a = self.bill.rows[3].cells[3].add_paragraph(str('' if item[2]=='0' else deciNum))
            Part6a.alignment = 2
            Part6b = self.bill.rows[3].cells[2].add_paragraph(str(item[1]))
            Part6b.alignment = 1
        self.bill.rows[3].cells[0].add_paragraph()
        self.bill.rows[3].cells[3].add_paragraph()
        self.bill.rows[3].cells[0].add_paragraph()
        self.bill.rows[3].cells[3].add_paragraph()
        part7 = self.bill.rows[3].cells[0].add_paragraph()
        part8 = self.bill.rows[3].cells[3].add_paragraph()
        part7.add_run('Add:').bold = True
        loop_counter = 0
        if len(GstRates)>1:
            for key in GstRates.keys():
                if loop_counter==0:
                    part7.add_run(f'\t{key}@ {GstRates[key][0]}%')
                    part8.alignment=2
                    part8.add_run(str(GstRates[key][1]))
                else:
                    self.bill.rows[3].cells[0].add_paragraph(f'\t{key}@ {GstRates[key][0]}%')
                    part9 = self.bill.rows[3].cells[3].add_paragraph()
                    part9.add_run(str(GstRates[key][1]))
                    part9.alignment =2
                loop_counter+=1
        else:
            for key in GstRates.keys():
                part7.add_run(f'\t{key}@ {GstRates[key][0]}%')
                part8.alignment=2
                GstAmount = format_decimal(float(GstRates[key][1]), locale='en_IN')
                decimNumber = GstAmount if GstAmount.find('.')!=-1 else GstAmount+'.00' 
                part8.add_run(str(decimNumber))
        Amount = self.bill.rows[4].cells[0].paragraphs[0]
        Amount.add_run(currency.currency((TotalAmount.replace(',','')))).bold = True
        Total = self.bill.rows[4].cells[2].paragraphs[0]
        Total.alignment = 1
        Total.add_run("Total").bold = True
        TotalAmountValue = self.bill.rows[4].cells[3].paragraphs[0]
        TotalAmountValue.alignment = 2
        TotalValue = format_decimal(float(TotalAmount), locale='en_IN')
        TotalAmountValue.add_run(TotalValue if TotalValue.find('.')!=-1 else TotalValue+'.00').bold = True
        self.document.add_paragraph()
        self.document.add_paragraph()

    def infoTable(self,infoList, BankList):
        Paninfo = self.document.add_table(2,2,"Table Grid")
        for cell in Paninfo.columns[0].cells:
                cell.width = Cm(2.1)
        for cell in Paninfo.columns[1].cells:
                cell.width = Cm(3.97)
        PanTitle = Paninfo.rows[0].cells[0].paragraphs[0]
        PanTitle.add_run('PAN:').bold=True
        PanTitleValue = Paninfo.rows[0].cells[1].paragraphs[0]
        PanTitleValue.add_run(infoList[0]).bold=True
        GSTTitle = Paninfo.rows[1].cells[0].paragraphs[0]
        GSTTitle.add_run('GSTIN:').bold=True
        GSTTitleValue = Paninfo.rows[1].cells[1].paragraphs[0]
        GSTTitleValue.add_run(infoList[1]).bold=True
        self.document.add_paragraph()
        titletable = self.document.add_table(1,1,"Table Grid")
        titletable.alignment = WD_TABLE_ALIGNMENT.RIGHT
        titletable.columns[0].cells[0].width = Cm(7.17)
        title = titletable.rows[0].cells[0].paragraphs[0]
        title.add_run('Bank Account Details for NEFT').bold=True
        self.document.add_paragraph()
        bankInfo = self.document.add_table(6,2,"Table Grid")
        for cell in bankInfo.columns[0].cells:
                cell.width = Cm(2.89)
        for cell in bankInfo.columns[1].cells:
                cell.width = Cm(4.44)
        bankInfo.alignment = WD_TABLE_ALIGNMENT.RIGHT
        NameTitle = bankInfo.rows[0].cells[0].paragraphs[0]
        NameTitle.add_run('Name:').bold=True
        Name = bankInfo.rows[0].cells[1].paragraphs[0]
        Name.add_run(BankList[0])

        AccountTitle = bankInfo.rows[1].cells[0].paragraphs[0]
        AccountTitle.add_run('Account No.:').bold=True
        Account = bankInfo.rows[1].cells[1].paragraphs[0]
        Account.add_run(BankList[1])

        IFSCTitle = bankInfo.rows[2].cells[0].paragraphs[0]
        IFSCTitle.add_run('IFSC Code :').bold=True
        IFSC = bankInfo.rows[2].cells[1].paragraphs[0]
        IFSC.add_run(BankList[2])

        MICRTitle = bankInfo.rows[3].cells[0].paragraphs[0]
        MICRTitle.add_run('MICR Code :').bold=True
        MICR = bankInfo.rows[3].cells[1].paragraphs[0]
        MICR.add_run(BankList[3])

        BankTitle = bankInfo.rows[4].cells[0].paragraphs[0]
        BankTitle.add_run('Bank:').bold=True
        Bank = bankInfo.rows[4].cells[1].paragraphs[0]
        Bank.add_run(BankList[4])

        BranchTitle = bankInfo.rows[5].cells[0].paragraphs[0]
        BranchTitle.add_run('Branch:').bold=True
        Branch = bankInfo.rows[5].cells[1].paragraphs[0]
        Branch.add_run(BankList[5])
        
        

    def saveDoc(self,filename):
        xfile=os.path.splitext(filename)
        if xfile[1]!='docx':
            filename = xfile[0]+'.docx'
        self.document.save(filename)

                    

