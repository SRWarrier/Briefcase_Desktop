from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn,nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
import docx
import datetime
import re
from num2words import num2words
import os
import time


def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def suffix(d):
            return 'th' if 11<=d<=13 else {1:'st',2:'nd',3:'rd'}.get(d%10, 'th')

def custom_strftime(format, t):
    return t.strftime(format).replace('{S}', str(t.day) + suffix(t.day))

class createDoc():
    def __init__(self):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Book Antiqua'
        font.size = docx.shared.Pt(11)

    def letterhead(self,companyName, reg_off,cin, email,phone,NoLetterHead = False):
        if not NoLetterHead:
            self.title = self.document.add_paragraph()
            self.title.alignment = 1
            self.title.paragraph_format.space_after = docx.shared.Pt(0)
            self.addressblock = self.document.add_paragraph()
            self.addressblock.paragraph_format.line_spacing = 1
            self.addressblock.alignment = 1
            self.addressblock.paragraph_format.space_after = 0
            self.addressblock.paragraph_format.space_before = 0
            COMP_NAME = self.title.add_run(companyName)
            COMP_NAME.bold = True
            COMP_NAME.font.size = docx.shared.Pt(14)
            REG_ADD=self.addressblock.add_run(f'Regd. Office: ')
            self.splitbysup(self.addressblock,reg_off,fontSize=10)
            REG_ADD.font.size = docx.shared.Pt(10)
            self.addressblock = self.document.add_paragraph()
            self.addressblock.paragraph_format.line_spacing = 1
            self.addressblock.alignment = 1
            self.addressblock.paragraph_format.space_after = 0
            self.addressblock.paragraph_format.space_before = 0
            CIN=self.addressblock.add_run(f'CIN: {cin}, ')
            CIN.font.size = docx.shared.Pt(10)
            emailid = self.addressblock.add_run(f'Email Id.: {email}, ')
            emailid.underline = True
            emailid.font.color.rgb = RGBColor(0, 0, 255)
            emailid.font.size = docx.shared.Pt(10)
            telno = self.addressblock.add_run(f'Tel. No.: {phone}')
            telno.font.size = docx.shared.Pt(10)
            insertHR(self.addressblock)
        else:
            self.document.add_paragraph("[ON THE LETTERHEAD OF THE COMPANY]")

    def Noticetitle2(self,MeetingType):
        if 'board' in MeetingType.lower():
            self.meetingtype = 'Meeting of the Board of Directors'
        elif 'extra' in MeetingType.lower():
            self.meetingtype  = 'Extra-Ordinary General Meeting'
        elif 'committee' in MeetingType.lower():
            self.meetingtype  = 'Meeting of the members of {committeeName}'
        else:
            self.meetingtype  = 'Annual General Meeting'
        self.NoticeTitle = self.document.add_paragraph()
        self.NoticeTitle.alignment = 1
        self.NoticeTitle.bold = True
        TitleText = self.NoticeTitle.add_run('NOTICE OF '+self.meetingtype.upper())
        TitleText.bold =True
        TitleText.underline =True

    def Noticetitle(self):
        self.document.add_paragraph()
        self.NoticeTitle = self.document.add_paragraph()
        self.NoticeTitle.alignment = 1
        self.NoticeTitle.bold = True
        TitleText = self.NoticeTitle.add_run('NOTICE')
        TitleText.bold =True
        TitleText.underline =True

    def NoticeTo(self, Addressees):
        loopCounter=1
        for addressee in  Addressees:
            if 'director' in addressee[2].lower():
                self.Addressee = self.document.add_paragraph()
                self.Addressee.paragraph_format.space_after = Pt(0)
                self.Addressee.paragraph_format.line_spacing = 1
                if addressee[1].lower()=='male':
                    title = 'Mr.'
                else:
                    title = 'Ms.'
                self.Addressee.add_run("\t"+str(loopCounter)+".\t"+title+addressee[0].title())
                loopCounter+=1

    def NoticeBody(self,MeetinNumber,meetingtype,MeetingName,finyear,TMZ,companyName, address, date, time,isRegOff=False,includeAgenda=True, Agenda=''):
        #mdate =  datetime.datetime.strptime(date,'%d/%m/%Y') 
        self.NoticeAddress = self.document.add_paragraph('')
        NoticeText = self.document.add_paragraph()
        NoticeText.alignment=3
        NoticeText.add_run("Notice is hereby given that ")
        if not MeetinNumber=='':
            MeetingNum = num2words(int(MeetinNumber),True)
            NoticeText.add_run(MeetingNum+' ')
        NoticeText.add_run(meetingtype)
        NoticeText.add_run(' of ')
        NoticeText.add_run(companyName).bold=True
        if MeetingName=='Board Meeting' or meetingtype=='Committee Meeting':
            NoticeText.add_run(f' for the financial year {finyear}')
        NoticeText.add_run(' will be held on ')
        self.splitbysup(NoticeText,custom_strftime('%A, {S} %B, %Y', date))
        NoticeText.add_run()
        NoticeText.add_run(' at ')  
        time = time.strftime(f'%I:%M %p ({TMZ}) ')
        NoticeText.add_run(time)
        NoticeText.add_run(f"{'at the Registered Office of the Company' if isRegOff else ''}")
        NoticeText.add_run(' at ')
        self.splitbysup(NoticeText,address)
        if includeAgenda:
            NoticeText.add_run(' for transacting the business mentioned in the Agenda;')
            AgendaTitle= self.document.add_paragraph()
            AgendaTitle.alignment = 1
            AgendaTitle.bold = True
            TitleText = AgendaTitle.add_run('AGENDA & NOTES')
            TitleText.underline =True
            TitleText.bold = True
            AgendaTable = self.document.add_table(len(Agenda)+1,3,style='Table Grid')
            AgendaHeaders = ['Sl No.','BUSINESS TO BE TRANSACTED','SUBMITTED TO BOARD FOR']
            for cell in range(len(AgendaTable.rows[0].cells)):
                Bg_color = parse_xml(r'<w:shd {} w:fill="e0e0e0"/>'.format(nsdecls('w')))
                AgendaTable.rows[0].cells[cell]._tc.get_or_add_tcPr().append(Bg_color)
            for x in range(3):
                HeadText = AgendaTable.rows[0].cells[x].paragraphs[0]
                HeadText.alignment = 1
                AgendaTable.rows[0].cells[x].width = Cm(15)
                Text = HeadText.add_run(AgendaHeaders[x])
                Text.font.bold = True
                for cell in AgendaTable.columns[0].cells:
                    cell.width = Cm(1)
                for cell in AgendaTable.columns[2].cells:
                    cell.width = Cm(4)
            for x in range(len(Agenda)):
               AgendaNo = AgendaTable.rows[x+1].cells[0].paragraphs[0]
               AgendaNo.alignment=3
               AgendaNo.paragraph_format.line_spacing = 1.15
               AgendaNoEl = AgendaNo.add_run(str(x+1)+'.')
               AgendaText = AgendaTable.rows[x+1].cells[1].paragraphs[0]
               AgendaText.alignment=3
               AgendaText.paragraph_format.line_spacing = 1.15
               self.splitbysup(AgendaText,Agenda[x][0])
               if Agenda[x][1]!='':
                   AgendaTable.rows[x+1].cells[1].add_paragraph()
                   AgendaText = AgendaTable.rows[x+1].cells[1].paragraphs[1]
                   AgendaText.alignment=3
                   AgendaText.paragraph_format.line_spacing = 1
                   self.splitbysup(AgendaText,Agenda[x][1],10)
               ActionText = AgendaTable.rows[x+1].cells[2].paragraphs[0]
               ActionText.paragraph_format.line_spacing = 1.15
               ActionText.alignment=3
               self.splitbysup(ActionText,Agenda[x][2])
               ActionText.alignment = 1
               AgendaTable.rows[x+1].cells[2].vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
               
    def NoticeSignatures(self, CompanyName, Name, Designation, Din, Address):
        forBlock = self.document.add_paragraph('\n\n')
        forBlock.alignment = 1
        forBlock.add_run('for ')
        forBlock.add_run(CompanyName).bold=True
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)
        NameBlock = self.document.add_paragraph(Name.title())
        NameBlock.paragraph_format.line_spacing = 1
        NameBlock.paragraph_format.space_after = Pt(0)
        NameBlock.alignment =1
        DesignationBlock=self.document.add_paragraph("("+Designation+", "+"DIN:"+Din+")")
        DesignationBlock.paragraph_format.line_spacing = 1
        DesignationBlock.alignment =1
        DesignationBlock.paragraph_format.space_after = Pt(0)
        AddressList = Address.split(',')
        AddressString = ''
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<40 and len(AddressString+', '+AddressList[x].strip())<40:
                AddressString = AddressString+', '+AddressList[x].strip()
            else:
                AddressM = self.document.add_paragraph()
                self.splitbysup(AddressM,AddressString)
                AddressM.paragraph_format.line_spacing = 1
                AddressM.alignment = 1
                AddressM.paragraph_format.space_after = Pt(0)
                AddressString = ''
                AddressString = AddressString+AddressList[x].strip()
        if AddressString!='':
            AddressM = self.document.add_paragraph()
            self.splitbysup(AddressM,AddressString)
            AddressM.paragraph_format.line_spacing = 1
            AddressM.alignment = 1
            AddressM.paragraph_format.space_after = Pt(0)

    def splitbysup(self,paragraph,text,fontSize = 11,isbold = False):
       splitTxt =  re.split(r'(\d(TH|Th|ST|St|th|st|nd|Nd|ND|rd|Rd|RD))',text)
       for x in splitTxt:
                if not(re.match('\d(TH|Th|ST|St|th|st|nd|Nd|ND|rd|Rd|RD)', x))==None:
                    text = paragraph.add_run(re.match('\d',x).group())
                    text.font.size = docx.shared.Pt(fontSize)
                elif x=='':
                    pass
                elif x.lower() in ('th','st','nd','rd'):
                    text = paragraph.add_run(x)
                    text.font.superscript = True
                    text.font.size = docx.shared.Pt(fontSize)
                else:
                    text = paragraph.add_run(x)
                    text.font.size = docx.shared.Pt(fontSize)
                if isbold:
                    text.font.bold = True
                    
    def saveDoc(self,filename):
        xfile=os.path.splitext(filename)
        if xfile[1]!='docx':
            filename = xfile[0]+'.docx'
        self.document.save(filename)
