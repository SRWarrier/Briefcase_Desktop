from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn,nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
import docx
import datetime
import re
from num2words import num2words
import os
import time



class createDoc():
    def __init__(self):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Book Antiqua'
        font.size = docx.shared.Pt(11)
    
    def HeadLine(self, head, doctitle, sectionpart):
        Bg_color = parse_xml(r'<w:shd {} w:fill="ebf2de"/>'.format(nsdecls('w')))
        border = self.document.add_table(1,1,style='Table Grid')
        border.rows[0].cells[0]._tc.get_or_add_tcPr().append(Bg_color)
        self.headline = border.rows[0].cells[0].add_paragraph()
        self.headline.paragraph_format.space_before=0
        self.headline.alignment = 1
        self.headline.paragraph_format.space_after = docx.shared.Pt(10)
        self.headline.paragraph_format.line_spacing = 1.125
        headText = self.headline.add_run(head)
        headText.font.bold = True
        headText.font.size = docx.shared.Pt(12)
        self.titleline = border.rows[0].cells[0].add_paragraph()
        self.titleline.paragraph_format.space_before=0
        self.titleline.alignment = 1
        self.titleline.paragraph_format.space_after = docx.shared.Pt(10)
        self.titleline.paragraph_format.line_spacing = 1.125
        titleText = self.titleline.add_run(doctitle)
        titleText.font.bold = True
        titleText.font.size = docx.shared.Pt(12)
        self.Sectionline = border.rows[0].cells[0].add_paragraph()
        self.Sectionline.paragraph_format.space_before=0
        self.Sectionline.alignment = 1
        self.Sectionline.paragraph_format.space_after = docx.shared.Pt(10)
        self.Sectionline.paragraph_format.line_spacing = 1.125
        SectionText = self.Sectionline.add_run(sectionpart)
        SectionText.font.size = docx.shared.Pt(11)
        SectionText.font.italic = True
    
    def NoticeAddressee(self, Name, Address,title=None):
        self.add_paragrpha_115s()
        self.add_paragrpha_115s('To,')
        if title:
            self.add_paragrpha_115s('{}'.format(title))
        Addressee = self.document.add_paragraph()
        Addressee.paragraph_format.line_spacing = 1.15
        Addressee.paragraph_format.space_after = Pt(0)
        AD = Addressee.add_run('{}'.format(Name))
        AD.bold = True
        AddressList = Address.split(',')
        AddressString = ''
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<50 and len(AddressString+', '+AddressList[x].strip())<50:
                AddressString = AddressString+', '+AddressList[x].strip()
                if x==len(AddressList)-1:
                    AddressM = self.document.add_paragraph(AddressString)
                    AddressM.paragraph_format.line_spacing = 1
                    AddressM.paragraph_format.space_after = Pt(0)
                    AddressString = ''
            else:
                AddressM = self.document.add_paragraph(AddressString)
                AddressM.paragraph_format.line_spacing = 1
                AddressM.paragraph_format.space_after = Pt(0)
                AddressString = ''
                AddressString = AddressString+AddressList[x].strip()
        if AddressString!='':
            AddressM = self.document.add_paragraph(AddressString)
            AddressM.paragraph_format.line_spacing = 1
            AddressM.paragraph_format.space_after = Pt(0)

    def add_paragrpha_115s(self,content=''):
        paragraph = self.document.add_paragraph(content)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(0)
        
        
    def MBP1(self, DirectorName, Gender, FathersName, DirectorAddress, AssocaitedCompaniesList, Designation, DIN, signdate, signplace):
        self.add_paragrpha_115s()
        paragraph = self.document.add_paragraph()
        DearSir = paragraph.add_run('Dear Sir(s)')
        paragraph2 = self.document.add_paragraph()
        text = paragraph2.add_run(f'I, {DirectorName}, {"son" if Gender.lower()=="male" else "daughter"} of {FathersName}, resident of {DirectorAddress}, being a director in the company hereby give notice of my interest or concern in the following company or companies, bodies corporate,firms or other association of individuals:-')
        paragraph2.alignment = 3
        MBPTable = self.document.add_table(len(AssocaitedCompaniesList)+1,5,style='Table Grid')
        MBPHeaders = ['Sl No.','Names of the Companies/bodies corporate/ firms/ association of individuals',
                      'Nature of interest or concern / Change in interest or concern','Shareholding',
                      'Date on which interest or concern arose / changed']
        for x in range(5):
            HeadText = MBPTable.rows[0].cells[x].paragraphs[0]
            MBPTable.rows[0].cells[x].width = Cm(5)
            Text = HeadText.add_run(MBPHeaders[x])
            for cell in MBPTable.columns[0].cells:
                cell.width = Cm(1)
        for x in range(len(AssocaitedCompaniesList)):
            for y in range(len(AssocaitedCompaniesList[x])):
               AssCompText = MBPTable.rows[x+1].cells[y].paragraphs[0]
               MBPTable.rows[x+1].cells[y].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
               AssCompText.paragraph_format.line_spacing = 1.125
               AssCompTextEl = AssCompText.add_run(str(AssocaitedCompaniesList[x][y]))
               AssCompText.alignment = 1
        for cell in MBPTable.columns[1].cells:
            cell.paragraphs[0].alignment = 3
               
        self.add_paragrpha_115s()
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('List of relatives is enclosed as ')
        paragraph.add_run('Annexure - A').font.bold = True
        paragraph.add_run(' to this notice.')
        self.add_paragrpha_115s()
        self.add_paragrpha_115s('Kindly ensure that this notice is brought upon and read at the ensuing meeting of the board.')
        self.add_paragrpha_115s()
        self.add_paragrpha_115s('Thanking you.')
        self.add_paragrpha_115s('Yours faithfully,')
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        NameBlock = self.document.add_paragraph("Name: "+DirectorName)
        NameBlock.paragraph_format.line_spacing = 1
        NameBlock.paragraph_format.space_after = Pt(0)
        DesignBlock = self.document.add_paragraph('(Designation: '+Designation+')')
        DesignBlock.paragraph_format.line_spacing = 1
        DesignBlock.paragraph_format.space_after = Pt(0)
        DINBlock = self.document.add_paragraph('(DIN: '+DIN+')')
        DINBlock.paragraph_format.line_spacing = 1
        DINBlock.paragraph_format.space_after = Pt(0)
        AddressList = DirectorAddress.split(',')
        AddressString = ''
        firstInsert = True
        FirstParagraph = self.document.add_paragraph()
        FirstParagraph.paragraph_format.line_spacing = 1
        FirstParagraph.paragraph_format.space_after = Pt(0)
        FirstParagraph.add_run('Address: ').font.bold = True
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<50 and len(AddressString+', '+AddressList[x].strip())<50:
                AddressString = AddressString+', '+AddressList[x].strip()
                if x==len(AddressList)-1:
                    AddressM = self.document.add_paragraph(AddressString)
                    AddressM.paragraph_format.line_spacing = 1
                    AddressM.paragraph_format.space_after = Pt(0)
                    AddressString = ''
            else:
                if firstInsert:
                    FirstParagraph.add_run(AddressString)
                    AddressString = ''
                    AddressString = AddressString+AddressList[x].strip()
                    firstInsert = False
                else:
                    AddressM = self.document.add_paragraph(AddressString)
                    AddressM.paragraph_format.line_spacing = 1
                    AddressM.paragraph_format.space_after = Pt(0)
                    AddressString = ''
                    AddressString = AddressString+AddressList[x].strip()
        if AddressString!='':
            AddressM = self.document.add_paragraph(AddressString)
            AddressM.paragraph_format.line_spacing = 1
            AddressM.paragraph_format.space_after = Pt(0)
        self.add_paragrpha_115s()
        self.add_paragrpha_115s()
        self.endLine(signdate,signplace)
            
            
    def endLine(self, signdate, signplace):
        Bg_color = parse_xml(r'<w:shd {} w:fill="ebf2de"/>'.format(nsdecls('w')))
        border = self.document.add_table(1,1,style='Table Grid')
        border.rows[0].cells[0]._tc.get_or_add_tcPr().append(Bg_color)
        self.headline = border.rows[0].cells[0].paragraphs[0]
        self.headline.paragraph_format.space_before=0
        self.headline.paragraph_format.space_after = docx.shared.Pt(0)
        self.headline.paragraph_format.line_spacing = 1.125
        headText = self.headline.add_run(f'Date: {signdate}')
        headText.font.size = docx.shared.Pt(11)
        self.titleline = border.rows[0].cells[0].add_paragraph()
        self.titleline.paragraph_format.space_before=0
        self.titleline.paragraph_format.space_after = docx.shared.Pt(0)
        self.titleline.paragraph_format.line_spacing = 1.125
        titleText = self.titleline.add_run(f'Place: {signplace}')
        titleText.font.size = docx.shared.Pt(11)
         
    def saveDoc(self,filename):
        xfile=os.path.splitext(filename)
        if xfile[1]!='docx':
            filename = xfile[0]+'.docx'
        self.document.save(filename)
