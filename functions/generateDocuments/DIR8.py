from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn,nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL,WD_TABLE_ALIGNMENT
import docx
import datetime
import re
from num2words import num2words
import os
import time



class createDoc():
    def __init__(self):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Book Antiqua'
        font.size = docx.shared.Pt(11)
    
    def HeadLine(self, head, doctitle, sectionpart):
        Bg_color = parse_xml(r'<w:shd {} w:fill="ebf2de"/>'.format(nsdecls('w')))
        border = self.document.add_table(1,1,style='Table Grid')
        border.rows[0].cells[0]._tc.get_or_add_tcPr().append(Bg_color)
        self.headline = border.rows[0].cells[0].add_paragraph()
        self.headline.paragraph_format.space_before=0
        self.headline.alignment = 1
        self.headline.paragraph_format.space_after = docx.shared.Pt(10)
        self.headline.paragraph_format.line_spacing = 1.125
        headText = self.headline.add_run(head)
        headText.font.bold = True
        headText.font.size = docx.shared.Pt(12)
        self.titleline = border.rows[0].cells[0].add_paragraph()
        self.titleline.paragraph_format.space_before=0
        self.titleline.alignment = 1
        self.titleline.paragraph_format.space_after = docx.shared.Pt(10)
        self.titleline.paragraph_format.line_spacing = 1.125
        titleText = self.titleline.add_run(doctitle)
        titleText.font.bold = True
        titleText.font.size = docx.shared.Pt(12)
        self.Sectionline = border.rows[0].cells[0].add_paragraph()
        self.Sectionline.paragraph_format.space_before=0
        self.Sectionline.alignment = 1
        self.Sectionline.paragraph_format.space_after = docx.shared.Pt(10)
        self.Sectionline.paragraph_format.line_spacing = 1.125
        SectionText = self.Sectionline.add_run(sectionpart)
        SectionText.font.size = docx.shared.Pt(11)
        SectionText.font.italic = True
    
    def NoticeAddressee(self, Name,title=None):
        self.add_paragrpha_115s()
        self.add_paragrpha_115s('To,')
        if title:
            self.add_paragrpha_115s('{}'.format(title))
        Addressee = self.document.add_paragraph()
        Addressee.paragraph_format.line_spacing = 1.15
        Addressee.paragraph_format.space_after = Pt(0)
        Addressee.add_run('Board of Directors of ')
        Addressee.add_run(f'“{Name}”').font.bold = True
       

    def add_paragrpha_115s(self,content=''):
        paragraph = self.document.add_paragraph(content)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(0)
        
    def preBlock(self,RegNo,NomCap,PaidCap, Name, Address):
        self.add_paragrpha_115s()
        preBlockTable = self.document.add_table(5,3)
        preBlockTable.alignment = WD_TABLE_ALIGNMENT.CENTER
        DataList = [['Registration No. of Company',':',RegNo],
                    ['Nominal Capital',':',NomCap],
                    ['Paid-up Capital',':',PaidCap],
                    ['Name of Company',':',Name],
                    ['Address of its Registered Office',':',Address]]
        for x in range(len(DataList)):
            for y in range(len(DataList[x])):
               AssCompText = preBlockTable.rows[x].cells[y].paragraphs[0]
               preBlockTable.rows[x].cells[y].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
               AssCompText.paragraph_format.line_spacing = 1.15
               AssCompTextEl = AssCompText.add_run(str(DataList[x][y]))
               AssCompText.alignment = 3
        for col in  range(3):
            if col==0:
                cellwid = 6.69
            elif col==1:
                cellwid = 0.5
            else:
                cellwid = 8.43
            for cell in preBlockTable.columns[col].cells:
                cell.width = Cm(cellwid)

    def DIR8(self, DirectorName, Gender, FathersName,DirectorAddress, AssocaitedCompaniesList, Designation, DIN, signdate, signplace):
        self.add_paragrpha_115s()
        paragraph2 = self.document.add_paragraph()
        text = paragraph2.add_run(f'I, {DirectorName}, {"son" if Gender.lower()=="male" else "daughter"} of {FathersName}, resident of {DirectorAddress}, Director in the Company, hereby give notice that I am/was a director in the following companies during the last three years: -')
        paragraph2.alignment = 3
        MBPTable = self.document.add_table(len(AssocaitedCompaniesList)+1,3,style='Table Grid')
        MBPTable.alignment = WD_TABLE_ALIGNMENT.CENTER
        MBPHeaders = ['Name of the Company','Date of Appointment','Date of Cessation']
        for x in range(3):
            HeadText = MBPTable.rows[0].cells[x].paragraphs[0]
            MBPTable.rows[0].cells[x].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            MBPTable.rows[0].cells[x].width = Cm(5)
            Text = HeadText.add_run(MBPHeaders[x])
        for x in range(len(AssocaitedCompaniesList)):
            for y in range(len(AssocaitedCompaniesList[x])):
               AssCompText = MBPTable.rows[x+1].cells[y].paragraphs[0]
               MBPTable.rows[x+1].cells[y].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
               AssCompText.paragraph_format.line_spacing = 1.15
               AssCompTextEl = AssCompText.add_run(str(AssocaitedCompaniesList[x][y]))
               AssCompText.alignment = 1
        for cell in MBPTable.columns[0].cells:
            cell.paragraphs[0].alignment = 3
        HeadText.alignment = 1
               
        self.add_paragrpha_115s()
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run('I further confirm that I have not incurred disqualification under section 164(2) of the Companies Act, 2013 in any of the above Companies, in the previous financial year, and that I, at present, stand free from any disqualification from being a Director.')
        self.document.add_paragraph().paragraph_format.space_after = Pt(10)
        self.document.add_paragraph().paragraph_format.space_after = Pt(10)
        NameBlock = self.document.add_paragraph(DirectorName)
        NameBlock.paragraph_format.space_after = Pt(0)
        NameBlock.alignment = 2
        spaceafter = (len(DirectorName)/2)
        DinBlock = self.document.add_paragraph('(DIN: '+DIN+')'+' '*((int(spaceafter))+(int(spaceafter/2))))
        DinBlock.alignment = 2
        DinBlock.paragraph_format.space_after = Pt(0)
        self.add_paragrpha_115s()
        self.add_paragrpha_115s()
        self.endLine(signdate,signplace)
            
            
    def endLine(self, signdate, signplace):
        Bg_color = parse_xml(r'<w:shd {} w:fill="ebf2de"/>'.format(nsdecls('w')))
        border = self.document.add_table(1,1,style='Table Grid')
        border.rows[0].cells[0]._tc.get_or_add_tcPr().append(Bg_color)
        self.headline = border.rows[0].cells[0].paragraphs[0]
        self.headline.paragraph_format.space_before=0
        self.headline.paragraph_format.space_after = docx.shared.Pt(0)
        self.headline.paragraph_format.line_spacing = 1.125
        headText = self.headline.add_run(f'Date: {signdate}')
        headText.font.size = docx.shared.Pt(11)
        self.titleline = border.rows[0].cells[0].add_paragraph()
        self.titleline.paragraph_format.space_before=0
        self.titleline.paragraph_format.space_after = docx.shared.Pt(0)
        self.titleline.paragraph_format.line_spacing = 1.125
        titleText = self.titleline.add_run(f'Place: {signplace}')
        titleText.font.size = docx.shared.Pt(11)
         
    def saveDoc(self,filename):
        xfile=os.path.splitext(filename)
        if xfile[1]!='docx':
            filename = xfile[0]+'.docx'
        self.document.save(filename)
