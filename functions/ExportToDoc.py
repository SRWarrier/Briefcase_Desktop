from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn,nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor, Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
import docx
import datetime
import re
from num2words import num2words
import os
import time

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def suffix(d):
            return 'th' if 11<=d<=13 else {1:'st',2:'nd',3:'rd'}.get(d%10, 'th')

def custom_strftime(format, t):
    return t.strftime(format).replace('{S}', str(t.day) + suffix(t.day))

class createDoc():
    def __init__(self):
        self.document = docx.Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Book Antiqua'
        font.size = docx.shared.Pt(11)
    
    def letterhead(self,companyName, reg_off,cin, email,phone):
        self.title = self.document.add_paragraph()
        self.title.alignment = 1
        self.title.paragraph_format.space_after = docx.shared.Pt(0)
        self.addressblock = self.document.add_paragraph()
        self.addressblock.paragraph_format.line_spacing = 1
        self.addressblock.alignment = 1
        COMP_NAME = self.title.add_run(companyName)
        COMP_NAME.bold = True
        COMP_NAME.font.size = docx.shared.Pt(14)
        REG_ADD=self.addressblock.add_run(f'Regd. Office: {reg_off}'+'\n')
        REG_ADD.font.size = docx.shared.Pt(10)
        CIN=self.addressblock.add_run(f'CIN: {cin}, ')
        CIN.font.size = docx.shared.Pt(10)
        emailid = self.addressblock.add_run(f'Email Id.: {email}, ')
        emailid.underline = True
        emailid.font.color.rgb = RGBColor(0, 0, 255)
        emailid.font.size = docx.shared.Pt(10)
        telno = self.addressblock.add_run(f'Tel. No.: {phone}')
        telno.font.size = docx.shared.Pt(10)
        insertHR(self.addressblock)

    def HeadLine(self, head, doctitle, sectionpart):
        Bg_color = parse_xml(r'<w:shd {} w:fill="b2ceaa"/>'.format(nsdecls('w')))
        border = self.document.add_table(1,1)
        border.rows[0].cells[0]._tc.get_or_add_tcPr().append(Bg_color)
        self.headline = border.rows[0].cells[0].paragraphs[0]
        self.headline.paragraph_format.space_before=0
        self.headline.alignment = 1
        self.headline.paragraph_format.space_after = docx.shared.Pt(0)
        self.headline.paragraph_format.line_spacing = 1.125
        headText = self.headline.add_run(head)
        headText.font.bold = True
        headText.font.size = docx.shared.Pt(11)
        self.titleline = border.rows[0].cells[0].add_paragraph()
        self.titleline.paragraph_format.space_before=0
        self.titleline.alignment = 1
        self.titleline.paragraph_format.space_after = docx.shared.Pt(0)
        self.titleline.paragraph_format.line_spacing = 1.125
        titleText = self.titleline.add_run(doctitle)
        titleText.font.bold = True
        titleText.font.size = docx.shared.Pt(11)
        self.Sectionline = border.rows[0].cells[0].add_paragraph()
        self.Sectionline.paragraph_format.space_before=0
        self.Sectionline.alignment = 1
        self.Sectionline.paragraph_format.space_after = docx.shared.Pt(0)
        self.Sectionline.paragraph_format.line_spacing = 1.125
        SectionText = self.Sectionline.add_run(sectionpart)
        SectionText.font.size = docx.shared.Pt(11)
    
        
               
    def ExtractHeader(self,MeetingType,companyName, address, date, time):
        self.extractHeaderblock = self.document.add_paragraph()
        self.extractHeaderblock.alignment=3
        #mdate =  datetime.datetime.strptime(date,'%d/%m/%Y')
        time = time.strftime('%I:%M %p (IST)')
        ExtractHeader = "CERTIFIED EXTRACTS OF THE MINUTES OF THE {} OF M/S {} HELD ON {} AT {} AT {} TO TRANSACT FOLLOWING BUSINESS;".format(MeetingType.upper(),companyName.upper(),address.upper(),custom_strftime('%A, {S} %B, %Y', date).upper(),time.upper())
        extractHead = self.extractHeaderblock.add_run(ExtractHeader)
        extractHead.bold = True

    def ExtractText (self,title='', Narration='', resolution=''):
        self.titleblock = self.document.add_paragraph()
        self.titleblock.alignment=3
        TITLE = self.titleblock.add_run(title.upper())
        TITLE.bold = True
        TITLE.underline = True
        TITLE.alignment = 3
        if Narration:
            self.document.add_paragraph(Narration).alignment = 3
        self.document.add_paragraph(resolution).alignment = 3
        
    def Signatures(self, AuthorisedDict):
        self.forBlock = self.document.add_paragraph()
        self.forBlock.alignment = 1
        self.forBlock.add_run('//CERTIFIED TRUE COPY//\n').bold=True
        self.forBlock.add_run('for '+ AuthorisedDict['companyName']).bold=True
        self.SignatureBlock = self.document.add_paragraph()
        self.forBlock = self.document.add_table(1,3)
        self.forBlock.autofit=True
        if len(AuthorisedDict['signatories'])==1:
            colList=[1]
        elif len(AuthorisedDict['signatories'])==2:
            colList=[0,2]
        elif len(AuthorisedDict['signatories'])==3:
            colList=[0,1,2]
        for x in range(len(AuthorisedDict['signatories'])):
            NameBlock = self.forBlock.rows[0].cells[colList[x]].add_paragraph()
            NameBlock.alignment =1
            NameBlock.paragraph_format.line_spacing = 1
            NameBlock.add_run(AuthorisedDict['signatories'][x]['Name'])
            NameBlock.add_run("\n("+AuthorisedDict['signatories'][x]['Designation']+', DIN:'+AuthorisedDict['signatories'][x]['DIN']+")")
            NameBlock.add_run('\n'+AuthorisedDict['signatories'][x]['Address'])

    def NoticeSignatures(self, CompanyName, Name, Designation, Din, Address):
        forBlock = self.document.add_paragraph('\n\n')
        forBlock.alignment = 1
        forBlock.add_run('for ')
        forBlock.add_run(CompanyName).bold=True
        SignatureBlock = self.document.add_paragraph()
        self.document.add_paragraph()
        NameBlock = self.document.add_paragraph(Name)
        NameBlock.paragraph_format.line_spacing = 1
        NameBlock.paragraph_format.space_after = Pt(0)
        NameBlock.alignment =1
        DesignationBlock=self.document.add_paragraph("("+Designation+", "+"DIN:"+Din+")")
        DesignationBlock.paragraph_format.line_spacing = 1
        DesignationBlock.alignment =1
        DesignationBlock.paragraph_format.space_after = Pt(0)
        AddressList = Address.split(',')
        AddressString = ''
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<32 and len(AddressString+', '+AddressList[x].strip())<32:
                AddressString = AddressString+', '+AddressList[x].strip()
            else:
                AddressM = self.document.add_paragraph(AddressString)
                AddressM.paragraph_format.line_spacing = 1
                AddressM.alignment = 1
                AddressM.paragraph_format.space_after = Pt(0)
                AddressString = ''
                AddressString = AddressString+AddressList[x].strip()
        if AddressString!='':
            AddressM = self.document.add_paragraph(AddressString)
            AddressM.paragraph_format.line_spacing = 1
            AddressM.alignment = 1
            AddressM.paragraph_format.space_after = Pt(0)
            
    def Noticetitle2(self,MeetingType):
        if 'board' in MeetingType.lower():
            self.meetingtype = 'Meeting of the Board of Directors'
        elif 'extra' in MeetingType.lower():
            self.meetingtype  = 'Extra-Ordinary General Meeting'
        else:
            self.meetingtype  = 'Annual General Meeting'
        self.NoticeTitle = self.document.add_paragraph()
        self.NoticeTitle.alignment = 1
        self.NoticeTitle.bold = True
        TitleText = self.NoticeTitle.add_run('NOTICE OF '+self.meetingtype.upper())
        TitleText.bold =True
        TitleText.underline =True

    def Noticetitle(self):
        self.NoticeTitle = self.document.add_paragraph()
        self.NoticeTitle.alignment = 1
        self.NoticeTitle.bold = True
        TitleText = self.NoticeTitle.add_run('NOTICE')
        TitleText.bold =True
        TitleText.underline =True

    def NoticeAddressee(self, Name, Address,title=None):
        self.Addressee = self.document.add_paragraph()
        self.Addressee.add_run('\nTo')
        if title:
            self.Addressee.add_run('\n{}'.format(title))
        AD = self.Addressee.add_run('\n{}'.format(Name))
        AD.bold = True
        addressList = Address.split(',')
        loop_counter = 0
        while loop_counter!=len(addressList):
            if loop_counter==len(addressList)-1:
                self.Addressee.add_run('{}.'.format(addressList[loop_counter].strip()))
            elif loop_counter%3==0:
                self.Addressee.add_run('\n{}, '.format(addressList[loop_counter].strip()))
            else:
                self.Addressee.add_run('{}, '.format(addressList[loop_counter].strip()))
            loop_counter+=1

    def NoticeTo(self, Addressees):
        loopCounter=1
        for addressee in  Addressees[1:]:
            self.Addressee = self.document.add_paragraph()
            self.Addressee.paragraph_format.space_after = Pt(0)
            self.Addressee.paragraph_format.line_spacing = 1
            self.Addressee.add_run("\t"+str(loopCounter)+"\t"+addressee)
            loopCounter+=1
        

    def NoticeBody(self,MeetinNumber,meetingtype, companyName, address, date, time,isRegOff=False,includeAgenda=True, Agenda=''):
        #mdate =  datetime.datetime.strptime(date,'%d/%m/%Y') 
        self.NoticeAddress = self.document.add_paragraph('')
        NoticeText = self.document.add_paragraph()
        NoticeText.alignment=3
        NoticeText.add_run("Notice is hereby given that ")
        if not MeetinNumber=='':
            MeetingNum = num2words(int(MeetinNumber),True)
            NoticeText.add_run(MeetingNum+' ')
        NoticeText.add_run(meetingtype)
        NoticeText.add_run(' of ')
        NoticeText.add_run(companyName).bold=True
        NoticeText.add_run(' will be held on ')
        NoticeText.add_run(custom_strftime('%A, {S} %B, %Y', date))
        NoticeText.add_run(' at ')  
        time = time.strftime('%I:%M %p (IST) ')
        NoticeText.add_run(time)
        NoticeText.add_run(f"{'at the Registered Office of the Company' if isRegOff else ''}")
        NoticeText.add_run(' at ')
        NoticeText.add_run(address)
        if includeAgenda:
            NoticeText.add_run(' for transacting the business mentioned in the Agenda;')
            AgendaTitle= self.document.add_paragraph()
            AgendaTitle.alignment = 1
            AgendaTitle.bold = True
            TitleText = AgendaTitle.add_run('AGENDA & NOTES')
            TitleText.underline =True
            TitleText.bold = True
            AgendaTable = self.document.add_table(len(Agenda)+1,3,style='Table Grid')
            AgendaHeaders = ['Sl No.','BUSINESS TO BE TRANSACTED','SUBMITTED TO BOARD FOR']
            for cell in range(len(AgendaTable.rows[0].cells)):
                Bg_color = parse_xml(r'<w:shd {} w:fill="e0e0e0"/>'.format(nsdecls('w')))
                AgendaTable.rows[0].cells[cell]._tc.get_or_add_tcPr().append(Bg_color)
            for x in range(3):
                HeadText = AgendaTable.rows[0].cells[x].paragraphs[0]
                HeadText.alignment = 1
                AgendaTable.rows[0].cells[x].width = Cm(15)
                Text = HeadText.add_run(AgendaHeaders[x])
                Text.font.bold = True
                for cell in AgendaTable.columns[0].cells:
                    cell.width = Cm(1)
                for cell in AgendaTable.columns[2].cells:
                    cell.width = Cm(4)
            print(Agenda)
            for x in range(len(Agenda)):
               AgendaNo = AgendaTable.rows[x+1].cells[0].paragraphs[0]
               AgendaNo.alignment=3
               AgendaNo.paragraph_format.line_spacing = 1.15
               AgendaNoEl = AgendaNo.add_run(str(x+1)+'.')
               AgendaText = AgendaTable.rows[x+1].cells[1].paragraphs[0]
               AgendaText.alignment=3
               AgendaText.paragraph_format.line_spacing = 1.15
               AgendaTextEl = AgendaText.add_run(Agenda[x][0])
               if Agenda[x][1]!='':
                   AgendaTable.rows[x+1].cells[1].add_paragraph()
                   AgendaText = AgendaTable.rows[x+1].cells[1].paragraphs[1]
                   AgendaText.alignment=3
                   AgendaText.paragraph_format.line_spacing = 1
                   AgendaTextE2 = AgendaText.add_run(Agenda[x][1])
                   AgendaTextE2.font.size = docx.shared.Pt(10)
               ActionText = AgendaTable.rows[x+1].cells[2].paragraphs[0]
               ActionText.paragraph_format.line_spacing = 1.15
               ActionText.alignment=3
               ActionTextE1 = ActionText.add_run(Agenda[x][1])
               ActionText.alignment = 1
               AgendaTable.rows[x+1].cells[2].vertical_alignment  = WD_ALIGN_VERTICAL.CENTER
               
        
    def MBP1(self, DirectorName, Gender, FathersName, DirectorAddress, AssocaitedCompaniesList, Designation, DIN, signdate, signplace):
        paragraph = self.document.add_paragraph()
        DearSir = paragraph.add_run('Dear Sir(s)')
        DearSir.bold = True
        paragraph2 = self.document.add_paragraph()
        text = paragraph2.add_run(f'I, {DirectorName}, {"son" if Gender.lower()=="male" else "daughter"} of {FathersName}, resident of {DirectorAddress}, being a director in the company hereby give notice of my interest or concern in the following company or companies, bodies corporate,firms or other association of individuals:-')
        paragraph2.alignment = 3
        MBPTable = self.document.add_table(len(AssocaitedCompaniesList)+1,5,style='Table Grid')
        MBPHeaders = ['Sl No.','Names of the Companies/bodies corporate/ firms/ association of individuals',
                      'Nature of interest or concern / Change in interest or concern','Shareholding',
                      'Date on which interest or concern arose / changed']
        for x in range(5):
            HeadText = MBPTable.rows[0].cells[x].paragraphs[0]
            MBPTable.rows[0].cells[x].width = Cm(5)
            Text = HeadText.add_run(MBPHeaders[x])
            Text.font.bold = True
            for cell in MBPTable.columns[0].cells:
                cell.width = Cm(1)
        for x in range(len(AssocaitedCompaniesList)):
            for y in range(len(AssocaitedCompaniesList[x])):
               AssCompText = MBPTable.rows[x+1].cells[y].paragraphs[0]
               AssCompText.paragraph_format.line_spacing = 1.125
               AssCompTextEl = AssCompText.add_run(str(AssocaitedCompaniesList[x][y]))
        signature = self.document.add_paragraph('\n\n\n')
        NameBlock = signature.add_run(DirectorName)
        DesignBlock = signature.add_run('\n('+Designation+', DIN: '+DIN+')\n')
        AddressList = DirectorAddress.split(',')
        AddressString = ''
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<32 and len(AddressString+', '+AddressList[x].strip())<32:
                AddressString = AddressString+', '+AddressList[x].strip()
                if x==len(AddressList)-1:
                    signature.add_run(AddressString+'\n')
            else:
                signature.add_run(AddressString+',\n')
                AddressString = ''
                AddressString = AddressString+AddressList[x].strip()
                
        dateplace = self.document.add_paragraph('\n\n')
        dateplace.add_run(f'Date: {signdate}\n')
        dateplace.add_run(f'Place: {signplace}')


    def DIR8(self, RegNo, Nominal, Paidup, CompName, RegAddress, DirectorName, Gender, FathersName, DirectorAddress, AssocaitedCompaniesList, Designation, DIN, signdate, signplace):
        Infoparagraph = self.document.add_paragraph()
        RegNoText = Infoparagraph.add_run('Registration No. of Company: '+str(RegNo))
        NomiText = Infoparagraph.add_run('Nominal Capital Rs.'+str(Nominal))
        PaidText = Infoparagraph.add_run('Paid-up Capital Rs.'+str(Paidup))
        NameText = Infoparagraph.add_run('Name of Company: '+str(CompName))
        AddressText = Infoparagraph.add_run('Address of its Registered Office\n'+str(RegAddress))
        
        paragraph = self.document.add_paragraph()
        ToText = paragraph.add_run('To\n')
        BoardText = paragraph.add_run('The Board of Directors of '+CompName+'\n')
        paragraph2 = self.document.add_paragraph()
        text = paragraph2.add_run(f'I, {DirectorName}, {"son" if Gender.lower()=="male" else "daughter"} of {FathersName}, resident of {DirectorAddress}, being a {Designation} in the company hereby give notice that I am/was a director in the following companies during the last three years:-')
        paragraph2.alignment = 3
        MBPTable = self.document.add_table(len(AssocaitedCompaniesList)+1,5,style='Table Grid')
        MBPHeaders = ['Name of the Company','Date of Appointment','Date of Cessation']
        for x in range(3):
            HeadText = MBPTable.rows[0].cells[x].paragraphs[0]
            MBPTable.rows[0].cells[x].width = Cm(5)
            Text = HeadText.add_run(MBPHeaders[x])
            Text.font.bold = True
            for cell in MBPTable.columns[0].cells:
                cell.width = Cm(1)
        for x in range(len(AssocaitedCompaniesList)):
            for y in range(len(AssocaitedCompaniesList[x])):
               AssCompText = MBPTable.rows[x+1].cells[y].paragraphs[0]
               AssCompText.paragraph_format.line_spacing = 1.125
               AssCompTextEl = AssCompText.add_run(str(AssocaitedCompaniesList[x][y]))
        paragraph3 = self.document.add_paragraph()
        text2 = paragraph2.add_run('I further confirm that I have not incurred disqualification under section 164(2) of the Companies Act, 2013 in any of the above companies, in the previous financial year, and that I, at present, stand free from any disqualification from being a director.')
        paragraph4 = self.document.add_paragraph()
        paragraph4.alignment = 1
        OrText = paragraph4.add_run('or')
        OrText.bold = True
        paragraph5 = self.document.add_paragraph()
        text3 = paragraph5.add_run('I further confirm that I have incurred disqualifications under section 164(2) of the Companies Act, 2013 in the following company(s) in the previous financial year, and that I, at present stand disqualified from being a director.')
        text3.alignment = 1
        MBP2Table = self.document.add_table(len(AssocaitedCompaniesList)+1,5,style='Table Grid')
        MBP2Headers = ['Name of the Company','Date of Appointment','Date of Cessation']
        for x in range(3):
            HeadText = MBP2Table.rows[0].cells[x].paragraphs[0]
            MBP2Table.rows[0].cells[x].width = Cm(5)
            Text = HeadText.add_run(MBP2Headers[x])
            Text.font.bold = True
            for cell in MBP2Table.columns[0].cells:
                cell.width = Cm(1)
        signature = self.document.add_paragraph('\n\n\n')
        NameBlock = signature.add_run(DirectorName)
        DesignBlock = signature.add_run('\n('+Designation+', DIN: '+DIN+')\n')
        AddressList = DirectorAddress.split(',')
        AddressString = ''
        for x in range(len(AddressList)):
            if x==0:
                AddressString = AddressString+AddressList[x].strip()
            elif len(AddressString)<32 and len(AddressString+', '+AddressList[x].strip())<32:
                AddressString = AddressString+', '+AddressList[x].strip()
                if x==len(AddressList)-1:
                    signature.add_run(AddressString+'\n')
            else:
                signature.add_run(AddressString+',\n')
                AddressString = ''
                AddressString = AddressString+AddressList[x].strip()
        dateplace = self.document.add_paragraph('\n\n')
        dateplace.add_run(f'Date: {signdate}\n')
        dateplace.add_run(f'Place: {signplace}')
                          

    def letterClosing(self,isfor =False,companyName='',signedby='',designation='',place='',date='', includeAgenda=False):
        self.closingText = self.document.add_paragraph('\n\n')
        if isfor:
            self.closingText.add_run(f'For {companyName}'+'\n'*4)
        self.closingText.add_run(f'{signedby}\n')
        self.closingText.add_run(f'({designation})\n\n')
        self.closingText.add_run('Date:\n')
        self.closingText.add_run('Place:')
        if includeAgenda:
            self.closingText.add_run('\n\nEncl. Agenda')

    def splitbysup(self,paragraph,text, isbold = False):
       splitTxt =  re.split(r'(\d(TH|ST|th|st|nd|ND|rd|RD))',text)
       for x in splitTxt:
                if not(re.match('\d(TH|ST|th|st|nd|ND|rd|RD)', x))==None:
                    text = paragraph.add_run(re.match('\d',x).group())
                elif x=='':
                    pass
                elif x.lower() in ('th','st','nd','rd'):
                    text = paragraph.add_run(x)
                    text.font.superscript = True
                else:
                    text = paragraph.add_run(x)
                if isbold:
                    text.font.bold = True
                    
    def saveDoc(self,filename):
        xfile=os.path.splitext(filename)
        if xfile[1]!='docx':
            filename = xfile[0]+'.docx'
        self.document.save(filename)


    
